# Neo Stats Hackathon
# Amazon Financial Report 

# Loading the libraries

import os
import io
import base64
import requests
import PyPDF2
import pandas as pd
import jsonz
from dotenv import load_dotenv
from openai import AzureOpenAI
import azure.ai.documentintelligence as docint
from azure.core.credentials import AzureKeyCredential
import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

load_dotenv()

class FinancialSummaryGenerator:
    def __init__(self):
        
        self.azure_openai_client = AzureOpenAI(
            api_key=os.environ.get("AZURE_OPENAI_API_KEY"),
            api_version="2024-02-01",
            azure_endpoint=os.environ.get("AZURE_OPENAI_ENDPOINT")
        )
        
        endpoint = os.environ.get("DOCUMENTINTELLIGENCE_ENDPOINT")
        key = os.environ.get("DOCUMENTINTELLIGENCE_API_KEY")
        
        self.document_client = docint.DocumentIntelligenceClient(
            endpoint=endpoint, 
            credential=AzureKeyCredential(key)
        )

        self.deployment_name = os.environ.get("DEPLOYMENT_NAME", "gpt-35-turbo")

    def extract_pdf_text(self, pdf_path):
 
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                for page in pdf_reader.pages:
                    full_text += page.extract_text() + "\n\n"
                if len(full_text.strip()) < 100:
                    raise ValueError("Extracted text is too short")
                
                return full_text
        
        except Exception as e:
            print(f"PDF text extraction error: {e}")
            
            try:
                with open(pdf_path, 'rb') as file:
                    poller = self.document_client.begin_analyze_document(
                        "prebuilt-document", file
                    )
                    result = poller.result()
                    return result.content
            
            except Exception as azure_error:
                print(f"Azure Document Intelligence error: {azure_error}")
                raise ValueError(f"Could not extract text from PDF: {azure_error}")

    def calculate_financial_ratios(self, financial_text):
   
        ratio_prompt = f"""
        Analyze the following financial text and extract/calculate these key financial ratios:
        1. Current Ratio
        2. Debt-to-Equity Ratio
        3. Return on Equity (ROE)
        4. Return on Assets (ROA)
        5. Gross Profit Margin
        6. Net Profit Margin
        7. Earnings Per Share (EPS)

        Provide the calculations in a structured JSON format with:
        - Ratio Name
        - Calculated Value
        - Formula Used
        - Brief Explanation

        Financial Text:
        {financial_text[:5000]}  # Limit input size
        """

        try:
            response = self.azure_openai_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "system", "content": "You are a financial analyst expert in extracting financial ratios."},
                    {"role": "user", "content": ratio_prompt}
                ],
                max_tokens=1000,
                temperature=0.2
            )

            return response.choices[0].message.content
        except Exception as e:
            print(f"Ratio calculation error: {e}")
            return json.dumps({
                "error": "Could not calculate ratios",
                "details": str(e)
            })

    def generate_performance_summary(self, financial_text):
       
        summary_prompt = f"""
        Analyze the following financial text and create a comprehensive performance summary:
        
        1. Overall Financial Performance
        2. Key Business Highlights
        3. Strategic Achievements
        4. Financial Health Indicators
        5. Year-over-Year Comparisons

        Limit: Approximately 500 words

        Financial Text:
        {financial_text[:5000]} 
        """
        
        try:
            response = self.azure_openai_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "system", "content": "You are a professional financial report summarizer."},
                    {"role": "user", "content": summary_prompt}
                ],
                max_tokens=600,
                temperature=0.3
            )

            return response.choices[0].message.content
        except Exception as e:
            return f"Error generating performance summary: {e}"

    def generate_risk_summary(self, financial_text):
        """
        Extract and summarize risk factors
        """
        risk_prompt = f"""
        From the following financial text, identify and summarize key risk factors:
        
        Risk Categories:
        1. Market Risks
        2. Operational Risks
        3. Financial Risks
        4. Strategic Risks
        5. Emerging Challenges

        Provide a concise, actionable summary of potential risks.

        Financial Text:
        {financial_text[:5000]}  # Limit input size
        """
        
        try:
            response = self.azure_openai_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "system", "content": "You are an expert in extracting and summarizing risk factors from financial reports."},
                    {"role": "user", "content": risk_prompt}
                ],
                max_tokens=400,
                temperature=0.2
            )

            return response.choices[0].message.content
        except Exception as e:
            return f"Error generating risk summary: {e}"

    def create_excel_output(self, financial_ratios, output_path):
        
        """
        Create Excel file with financial ratios
        """
        try:
            if isinstance(financial_ratios, str):
                try:
                    ratios_dict = json.loads(financial_ratios)
                except json.JSONDecodeError:
      
                    ratios_dict = {"Raw_Output": financial_ratios}
            else:
                ratios_dict = financial_ratios

            df = pd.DataFrame.from_dict(ratios_dict, orient='index')
            df.to_excel(output_path, index=True)
            print(f"Excel output saved to {output_path}")
        except Exception as e:
            print(f"Excel creation error: {e}")
            
            error_df = pd.DataFrame([{"Error": str(e)}])
            error_df.to_excel(output_path)

    def create_word_document(self, performance_summary, risk_summary, output_path):

        try:
            document = Document()
            
      
            document.add_heading('Performance Summary', 0)
            performance_para = document.add_paragraph(performance_summary)
            performance_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            document.add_heading('Risk Factors Summary', 1)
            risk_para = document.add_paragraph(risk_summary)
            risk_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            document.save(output_path)
            print(f"Word document saved to {output_path}")
        except Exception as e:
            print(f"Word document creation error: {e}")

    def process_financial_report(self, pdf_path):
       
        try:
  
            financial_text = self.extract_pdf_text(pdf_path)
            
            financial_ratios = self.calculate_financial_ratios(financial_text)
    
            performance_summary = self.generate_performance_summary(financial_text)
            
            risk_summary = self.generate_risk_summary(financial_text)
            
            participant_name = os.environ.get("PARTICIPANT_NAME", "badri")
            excel_output_path = f'Neo_GenAI_hackathon_excel_{participant_name}.xlsx'
            self.create_excel_output(financial_ratios, excel_output_path)
            
            word_output_path = f'Neo_GenAI_hackathon_word_{participant_name}.docx'
            self.create_word_document(performance_summary, risk_summary, word_output_path)
            
            return {
                'excel_path': excel_output_path,
                'word_path': word_output_path
            }
        
        except Exception as e:
            print(f"Financial report processing error: {e}")
            return {"error": str(e)}

def main():
    required_vars = [
        "AZURE_OPENAI_API_KEY", 
        "AZURE_OPENAI_ENDPOINT", 
        "DOCUMENTINTELLIGENCE_ENDPOINT", 
        "DOCUMENTINTELLIGENCE_API_KEY"
    ]
    
    for var in required_vars:
        if not os.environ.get(var):
            print(f"Error: {var} environment variable is not set")
            return
    
    pdf_path = 'C:\\Users\\USER\\OneDrive\\Desktop\\NeoStats Hackathon\\Amazon-com-Inc-2023-Annual-Report.pdf'

    if not os.path.exists(pdf_path):
        print(f"Error: PDF file {pdf_path} not found")
        return
    
    generator = FinancialSummaryGenerator()
    outputs = generator.process_financial_report(pdf_path)
    
    if 'error' in outputs:
        print(f"Error: {outputs['error']}")
    else:
        print(f"Generated Excel file: {outputs['excel_path']}")
        print(f"Generated Word document: {outputs['word_path']}")

if __name__ == "__main__":
    main()